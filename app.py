import os
os.environ["ANONYMIZED_TELEMETRY"] = "False"
from datetime import datetime
import json
import re
import traceback
import numpy as np
# Patch NumPy 2.x compatibility for ChromaDB
if not hasattr(np, 'float_'):
    np.float_ = np.float64   # type: ignore
if not hasattr(np, 'int_'):  
    np.int_ = np.int64
if not hasattr(np, 'NaN'):
    np.NaN = np.nan          # type: ignore
import logging
logging.getLogger('chromadb').setLevel(logging.CRITICAL)
from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename

from database import db, init_db, CorpusDocument, AnalysisDocument, Detection, Setting
import vector_db
import embeddings

app = Flask(__name__, template_folder='templates', static_folder='static')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload size

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize database
init_db(app)

# Helper functions for Settings
def get_setting(key, default):
    with app.app_context():
        setting = Setting.query.get(key)
        if setting:
            return setting.value
        return default

def setting_to_bool(value, default=False):
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    normalized = str(value).strip().lower()
    if normalized in {"1", "true", "yes", "on", "tak"}:
        return True
    if normalized in {"0", "false", "no", "off", "nie"}:
        return False
    return default

def set_setting(key, value):
    setting = Setting.query.get(key)
    if setting:
        setting.value = str(value)
    else:
        setting = Setting(key=key, value=str(value)) # type: ignore
        db.session.add(setting)
    db.session.commit()

# Document text extractors
def extract_text(file_path, filename):
    ext = os.path.splitext(filename.lower())[1]
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    elif ext in ['.docx', '.doc']:
        import docx
        doc = docx.Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return "\n".join(text)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

# Chunker for corpus documents
def chunk_text(text, chunk_size, overlap):
    # Simple word splitting
    words = text.split()
    chunks = []
    
    if len(words) <= chunk_size:
        return [" ".join(words)]
        
    step = chunk_size - overlap
    if step <= 0:
        step = 1
        
    for i in range(0, len(words), step):
        chunk_words = words[i:i + chunk_size]
        chunks.append(" ".join(chunk_words))
        if i + chunk_size >= len(words):
            break
            
    return chunks

# Character-aware sliding window segmenter for analyzed texts
def segment_text_with_indices(text, segment_size, overlap):
    # Find all words and their indices
    words_matches = list(re.finditer(r'\S+', text))
    if not words_matches:
        return []
        
    segments = []
    num_words = len(words_matches)
    
    if num_words <= segment_size:
        start_char = words_matches[0].start()
        end_char = words_matches[-1].end()
        segments.append({
            'text': text[start_char:end_char],
            'start_char': start_char,
            'end_char': end_char
        })
        return segments
        
    step = segment_size - overlap
    if step <= 0:
        step = 1
        
    for i in range(0, num_words, step):
        end_idx = min(i + segment_size, num_words)
        
        start_char = words_matches[i].start()
        end_char = words_matches[end_idx - 1].end()
        
        segments.append({
            'text': text[start_char:end_char],
            'start_char': start_char,
            'end_char': end_char
        })
        
        if end_idx >= num_words:
            break
            
    return segments

# Main View
@app.route('/')
def index():
    return render_template('index.html')

# Settings API
@app.route('/api/settings', methods=['GET', 'POST'])
def api_settings():
    if request.method == 'GET':
        settings = Setting.query.all()
        return jsonify({s.key: s.value for s in settings})
    else:
        data = request.json or {}
        for k, v in data.items():
            set_setting(k, v)
        return jsonify({"status": "success", "message": "Settings updated successfully"})

# Corpus API
@app.route('/api/corpus', methods=['GET', 'POST'])
def api_corpus():
    if request.method == 'GET':
        docs = CorpusDocument.query.order_by(CorpusDocument.added_at.desc()).all()
        result = []
        for doc in docs:
            # Let's count chunks in ChromaDB for this document
            try:
                collection = vector_db.get_collection()
                # ChromaDB get count by filtering metadata
                doc_chunks = collection.get(where={"corpus_doc_id": doc.id})
                chunk_count = len(doc_chunks["ids"]) if doc_chunks else 0
            except Exception:
                chunk_count = 0
                
            result.append({
                "id": doc.id,
                "title": doc.title,
                "author": doc.author or "Unknown",
                "added_at": doc.added_at.strftime("%Y-%m-%d %H:%M:%S"),
                "chunk_count": chunk_count
            })
        return jsonify(result)
        
    else:
        # Check if file is uploaded
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
            
        file = request.files['file']
        title = request.form.get('title', '').strip()
        author = request.form.get('author', '').strip()
        
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
            
        if not title:
            title = os.path.splitext(file.filename)[0] # type: ignore
             
        filename = secure_filename(file.filename)      # type: ignore
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            # Extract text
            text = extract_text(file_path, filename)
            if not text.strip():
                return jsonify({"error": "File is empty or could not be read"}), 400
                
            # Create CorpusDocument in SQLite
            doc = CorpusDocument(title=title, author=author if author else None, content=text)  # type: ignore
            db.session.add(doc)
            db.session.commit() # Commit to get doc.id
            
            # Chunking parameters from settings
            chunk_size = int(get_setting('chunk_size', '50'))
            overlap = int(get_setting('chunk_overlap', '10'))
            
            chunks = chunk_text(text, chunk_size, overlap)
            
            # Add to ChromaDB vector store
            vector_db.add_document_chunks(doc.id, chunks)
            
            # Clean up uploaded file
            os.remove(file_path)
            
            return jsonify({
                "status": "success",
                "message": "Document added to corpus successfully",
                "document": {
                    "id": doc.id,
                    "title": doc.title,
                    "author": doc.author or "Unknown",
                    "chunk_count": len(chunks)
                }
            })
            
        except Exception as e:
            db.session.rollback()
            if os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500

@app.route('/api/corpus/<int:doc_id>', methods=['PUT'])
def api_update_corpus_doc(doc_id):
    doc = CorpusDocument.query.get(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404
        
    data = request.json or {}
    title = data.get('title', '').strip()
    author = data.get('author', '').strip()
    
    if not title:
        return jsonify({"error": "Title is required"}), 400
        
    try:
        doc.title = title
        doc.author = author if author else None
        db.session.commit()
        return jsonify({"status": "success", "message": "Document updated successfully"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/corpus/<int:doc_id>', methods=['DELETE'])
def api_delete_corpus_doc(doc_id):
    doc = CorpusDocument.query.get(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404
        
    try:
        # Delete from ChromaDB
        vector_db.delete_document_chunks(doc_id)
        
        # Delete from SQLite
        db.session.delete(doc)
        db.session.commit()
        
        return jsonify({"status": "success", "message": f"Document {doc_id} deleted successfully"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

# Helper function to merge overlapping texts word-by-word
def merge_overlapping_texts(text1, text2):
    if not text1:
        return text2
    if not text2:
        return text1
        
    words1 = text1.split()
    words2 = text2.split()
    
    max_check = min(len(words1), len(words2))
    
    # Find the maximum overlap at the end of words1 and start of words2
    overlap_size = 0
    for k in range(max_check, 0, -1):
        if words1[-k:] == words2[:k]:
            overlap_size = k
            break
            
    if overlap_size > 0:
        merged_words = words1 + words2[overlap_size:]
        return " ".join(merged_words)
    else:
        return text1 + " ... " + text2

def get_detection_source_fragments(det):
    fragments = getattr(det, '_source_fragments', None)
    if fragments is not None:
        return fragments

    chunk_idx = getattr(det, '_corpus_chunk_index', None)
    fragments = [{
        "chunk_index": chunk_idx,
        "text": det.corpus_snippet_text,
        "similarity": det.similarity_score,
        "match_type": det.match_type
    }]
    det._source_fragments = fragments
    return fragments

def dedupe_and_sort_source_fragments(fragments):
    seen = set()
    unique = []
    for fragment in sorted(
        fragments,
        key=lambda item: len((item.get("text") or "").strip()),
        reverse=True
    ):
        text = (fragment.get("text") or "").strip()
        chunk_index = fragment.get("chunk_index")
        key = (chunk_index, text)
        if not text or key in seen:
            continue
        if any(
            existing.get("chunk_index") == chunk_index and
            text.lower() in (existing.get("text") or "").lower()
            for existing in unique
        ):
            continue
        seen.add(key)
        unique.append(fragment)

    return sorted(
        unique,
        key=lambda fragment: (
            fragment.get("chunk_index") is None,
            fragment.get("chunk_index") if fragment.get("chunk_index") is not None else 0
        )
    )

def apply_source_fragments(det):
    fragments = dedupe_and_sort_source_fragments(get_detection_source_fragments(det))
    det._source_fragments = fragments
    det.corpus_source_fragments = json.dumps(fragments, ensure_ascii=False)
    det.corpus_snippet_text = "\n\n".join(
        f"[chunk {fragment.get('chunk_index')}] {fragment.get('text')}"
        for fragment in fragments
    )

# Helper function to merge close detections
def merge_close_detections(detections, input_text, max_gap_chars=150, max_chunk_gap=2):
    if not detections:
        return []

    def get_chunk_range(det):
        start = getattr(det, '_corpus_chunk_start', None)
        end = getattr(det, '_corpus_chunk_end', None)
        chunk_idx = getattr(det, '_corpus_chunk_index', None)

        if start is None:
            start = chunk_idx
        if end is None:
            end = chunk_idx

        return start, end

    def corpus_chunks_are_close(current_det, next_det):
        current_start, current_end = get_chunk_range(current_det)
        next_start, next_end = get_chunk_range(next_det)

        if None in (current_start, current_end, next_start, next_end):
            return False

        return (
            next_start <= current_end + max_chunk_gap and   # type: ignore
            next_end >= current_start - max_chunk_gap       # type: ignore
        )
        
    # Sort detections: primary by corpus_doc_id, then source position, then corpus chunk.
    sorted_dets = sorted(
        detections,
        key=lambda d: (
            d.corpus_document_id,
            d.start_char_idx,
            getattr(d, '_corpus_chunk_index', -1)
        )
    )
    
    merged = []
    current = sorted_dets[0]
    
    for next_det in sorted_dets[1:]:
        # If they point to the same document, are close in the analyzed text,
        # and are also close in the reference corpus, treat them as one passage.
        if (next_det.corpus_document_id == current.corpus_document_id and 
            next_det.start_char_idx <= current.end_char_idx + max_gap_chars and
            corpus_chunks_are_close(current, next_det)):
            
            # Merge!
            current.end_char_idx = max(current.end_char_idx, next_det.end_char_idx)
            current.source_snippet = input_text[current.start_char_idx:current.end_char_idx]
            current_chunk_start, current_chunk_end = get_chunk_range(current)
            next_chunk_start, next_chunk_end = get_chunk_range(next_det)
            current._corpus_chunk_start = min(current_chunk_start, next_chunk_start)  # type: ignore
            current._corpus_chunk_end = max(current_chunk_end, next_chunk_end)        # type: ignore
            
            # Keep max similarity score
            if next_det.similarity_score and current.similarity_score:
                current.similarity_score = max(current.similarity_score, next_det.similarity_score)
            elif next_det.similarity_score:
                current.similarity_score = next_det.similarity_score
                
            # Keep highest match severity: Full > Partial > Paraphrase > Allusion
            severity = {"Full": 4, "Partial": 3, "Paraphrase": 2, "Allusion": 1, "None": 0}
            curr_sev = severity.get(current.match_type, 0)
            next_sev = severity.get(next_det.match_type, 0)
            if next_sev > curr_sev:
                current.match_type = next_det.match_type
                
            # Combine explanations
            if next_det.explanation and next_det.explanation not in current.explanation:
                if current.explanation:
                    current.explanation = f"{current.explanation}\n[Ciąg dalszy]: {next_det.explanation}"
                else:
                    current.explanation = next_det.explanation
                
            current._source_fragments = dedupe_and_sort_source_fragments(
                get_detection_source_fragments(current) + get_detection_source_fragments(next_det)
            )
            apply_source_fragments(current)
        else:
            apply_source_fragments(current)
            merged.append(current)
            current = next_det
            
    apply_source_fragments(current)
    merged.append(current)
    return merged

def get_model_rates(model_name):
    """Returns standard Gemini API rates per 1M tokens in USD."""
    if "gemini-3.1-flash-lite" in model_name:
        return 0.25, 1.50
    return 1.50, 9.00

def estimate_api_cost(prompt_tokens, candidates_tokens, thoughts_tokens, embedding_tokens, model_name, provider="gemini"):
    embedding_cost = (embedding_tokens * 0.15) / 1000000.0
    if provider == "ollama":
        return embedding_cost

    input_rate, output_rate = get_model_rates(model_name)
    billable_output_tokens = candidates_tokens + thoughts_tokens

    prompt_cost = (prompt_tokens * input_rate) / 1000000.0
    output_cost = (billable_output_tokens * output_rate) / 1000000.0
    return prompt_cost + output_cost + embedding_cost

def add_usage_totals(usage, log_msg=None, context=""):
    if not usage:
        if log_msg:
            log_msg(f"[LLM] Brak usage_metadata dla wywołania: {context}")
        return 0, 0, 0, 1

    prompt_tokens = usage.get('prompt_tokens', 0)
    candidates_tokens = usage.get('candidates_tokens', 0)
    thoughts_tokens = usage.get('thoughts_tokens', 0)
    total_tokens = usage.get('total_tokens', 0)
    tool_use_prompt_tokens = usage.get('tool_use_prompt_tokens', 0)
    cached_content_tokens = usage.get('cached_content_tokens', 0)

    if log_msg:
        details = (
            f"prompt={prompt_tokens}, output={candidates_tokens}, "
            f"thinking={thoughts_tokens}, total={total_tokens}"
        )
        if tool_use_prompt_tokens or cached_content_tokens:
            details += f", tool_prompt={tool_use_prompt_tokens}, cached={cached_content_tokens}"
        log_msg(f"[LLM] Tokeny {context}: {details}")

    return prompt_tokens, candidates_tokens, thoughts_tokens, 1

def parse_source_fragments(det):
    if det.corpus_source_fragments:
        try:
            fragments = json.loads(det.corpus_source_fragments)
            if isinstance(fragments, list):
                return dedupe_and_sort_source_fragments(fragments)
        except json.JSONDecodeError:
            pass

    return dedupe_and_sort_source_fragments([{
        "chunk_index": None,
        "text": det.corpus_snippet_text,
        "similarity": det.similarity_score,
        "match_type": det.match_type
    }])

def parse_analysis_params(analysis):
    if analysis.analysis_params:
        try:
            params = json.loads(analysis.analysis_params)
            if isinstance(params, dict):
                return params
        except json.JSONDecodeError:
            pass
    return {}

def format_analysis_params(params):
    if not params:
        return "brak danych"

    parts = []
    provider = params.get("llm_provider")
    if provider:
        parts.append(f"provider: {provider}")
    if params.get("llm_model"):
        parts.append(f"model LLM: {params['llm_model']}")
    elif provider == "ollama" and params.get("ollama_model"):
        parts.append(f"model LLM: {params['ollama_model']}")
    elif provider == "gemini" and params.get("gemini_model"):
        parts.append(f"model LLM: {params['gemini_model']}")
    elif params.get("ollama_model"):
        parts.append(f"model LLM: {params['ollama_model']}")
    elif params.get("gemini_model"):
        parts.append(f"model LLM: {params['gemini_model']}")
    if params.get("embedding_model"):
        parts.append(f"embeddingi: {params['embedding_model']}")
    if params.get("ollama_base_url"):
        parts.append(f"Ollama URL: {params['ollama_base_url']}")
    if provider == "ollama" and params.get("ollama_think") is not None:
        parts.append(f"Ollama thinking: {'włączone' if params.get('ollama_think') else 'wyłączone'}")
    if params.get("gemini_temp") is not None:
        parts.append(f"temperatura: {float(params['gemini_temp']):.2f}")
    if provider != "ollama" and params.get("gemini_thinking_level"):
        parts.append(f"myślenie: {params['gemini_thinking_level']}")
    if params.get("top_k") is not None:
        parts.append(f"top_k: {params['top_k']}")
    if params.get("similarity_threshold") is not None:
        parts.append(f"próg podobieństwa: {float(params['similarity_threshold']):.2f}")
    if params.get("analysis_segment_size") is not None and params.get("analysis_segment_overlap") is not None:
        parts.append(
            f"segment analizy: {params['analysis_segment_size']}/{params['analysis_segment_overlap']} słów"
        )
    if params.get("corpus_chunk_size") is not None:
        parts.append(f"chunk korpusu: {params['corpus_chunk_size']} słów")
    if params.get("segments_count") is not None:
        parts.append(f"liczba segmentów: {params['segments_count']}")

    return "; ".join(parts) if parts else "brak danych"

def get_text_window_by_char(text, start_idx, end_idx, word_radius=8):
    words = list(re.finditer(r'\S+', text))
    if not words:
        return "", text[start_idx:end_idx], ""

    first_word_idx = 0
    last_word_idx = len(words) - 1
    for idx, match in enumerate(words):
        if match.end() > start_idx:
            first_word_idx = idx
            break
    for idx, match in enumerate(words):
        if match.start() < end_idx:
            last_word_idx = idx
        else:
            break

    context_start_word = max(0, first_word_idx - word_radius)
    context_end_word = min(len(words) - 1, last_word_idx + word_radius)
    context_start = words[context_start_word].start()
    context_end = words[context_end_word].end()

    return (
        text[context_start:start_idx],
        text[start_idx:end_idx],
        text[end_idx:context_end]
    )

def split_text_for_highlight(text, highlight):
    if not text:
        return "", "", ""
    if not highlight:
        return "", text, ""

    idx = text.find(highlight)
    if idx < 0:
        idx = text.lower().find(highlight.lower())
    if idx < 0:
        return text, "", ""

    end_idx = idx + len(highlight)
    return text[:idx], text[idx:end_idx], text[end_idx:]

def text_contains_highlight(text, highlight):
    if not text or not highlight:
        return False
    return bool(split_text_for_highlight(text, highlight)[1])

def add_shaded_run(paragraph, text, fill_color=None, bold=False):
    run = paragraph.add_run(text)
    run.bold = bold
    if fill_color:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>'
        run._r.get_or_add_rPr().append(parse_xml(shading_xml))
    return run

def add_highlighted_text_paragraph(doc, before, highlight, after, fill_color):
    paragraph = doc.add_paragraph()
    if before:
        paragraph.add_run(before)
    if highlight:
        add_shaded_run(paragraph, highlight, fill_color=fill_color, bold=False)
    if after:
        paragraph.add_run(after)
    return paragraph

def add_horizontal_rule(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BFC7D5')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return paragraph

def add_export_metadata(doc, analysis, detections_count=None):
    analysis_params = parse_analysis_params(analysis)
    if detections_count is None:
        detections_count = len(filter_visible_detections(analysis.detections))
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Data analizy: {analysis.analyzed_at.strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    meta_p.add_run(f"Liczba wykrytych zapożyczeń: {detections_count}\n").italic = True
    meta_p.add_run(
        f"LLM: {analysis.llm_calls or 0} wywołań, "
        f"tokeny: {analysis.prompt_tokens or 0} wej., "
        f"{analysis.candidates_tokens or 0} wyj., "
        f"{analysis.thoughts_tokens or 0} myśl.\n"
    ).italic = True
    meta_p.add_run(
        f"Szacowany koszt API: ${analysis.estimated_cost or 0.0:.6f}\n"
    ).italic = True
    meta_p.add_run(
        f"Parametry analizy: {format_analysis_params(analysis_params)}\n"
    ).italic = True

def safe_docx_filename(title, prefix):
    safe_title = "".join([c if c.isalnum() else "_" for c in title])
    return f"{prefix}_{safe_title}.docx"

def filter_visible_detections(detections):
    """Matches the frontend rule: hide detections overlapped by an earlier visible one."""
    sorted_detections = sorted(
        detections,
        key=lambda det: (
            det.start_char_idx,
            -(det.similarity_score or 0)
        )
    )

    visible = []
    last_end = 0
    for det in sorted_detections:
        if det.start_char_idx >= last_end:
            visible.append(det)
            last_end = det.end_char_idx

    return visible

# Analysis API
@app.route('/api/analyze', methods=['GET', 'POST'])
def api_analyze():
    if request.method == 'GET':
        analyses = AnalysisDocument.query.order_by(AnalysisDocument.analyzed_at.desc()).all()
        result = []
        for a in analyses:
            result.append({
                "id": a.id,
                "title": a.title,
                "analyzed_at": a.analyzed_at.strftime("%Y-%m-%d %H:%M:%S")
            })
        return jsonify(result)
        
    else:
        # Perform analysis
        title = request.form.get('title', '').strip()
        input_text = request.form.get('text', '').strip()
        
        # Check if file was uploaded instead of typing text
        if 'file' in request.files:
            file = request.files['file']
            if file.filename != '':
                filename = secure_filename(file.filename)   # type: ignore
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                try:
                    input_text = extract_text(file_path, filename)
                    if not title:
                        title = os.path.splitext(file.filename)[0]   # type: ignore
                except Exception as e:
                    return jsonify({"error": f"Failed to read file: {str(e)}"}), 500
                finally:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        
        if not input_text:
            return jsonify({"error": "No text provided for analysis"}), 400
            
        if not title:
            title = f"Analysis {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            
        # Initialize logs
        log_lines = []
        def log_msg(msg):
            timestamp = datetime.now().strftime('%H:%M:%S')
            line = f"[{timestamp}] {msg}"
            log_lines.append(line)
            print(f"[AnalysisLog] {line}", flush=True)

        log_msg("Rozpoczęto analizę tekstu...")
        log_msg(f"Tytuł analizy: {title}")
        log_msg(f"Długość tekstu: {len(input_text)} znaków")
        
        try:
            # 1. Save AnalysisDocument in SQLite
            analysis_doc = AnalysisDocument(title=title, content=input_text)   # type: ignore
            db.session.add(analysis_doc)
            db.session.commit() # Commit to get analysis_doc.id
            log_msg(f"Utworzono rekord analizy w bazie SQLite (ID: {analysis_doc.id})")
        except Exception as e:
            db.session.rollback()
            error_trace = traceback.format_exc()
            print(f"Błąd inicjalizacji analizy w bazie danych: {error_trace}")
            return jsonify({"error": f"Failed to initialize analysis in database: {str(e)}", "trace": error_trace}), 500
            
        temp_detections = []
        total_prompt_tokens = 0
        total_candidates_tokens = 0
        total_thoughts_tokens = 0
        total_embedding_tokens = 0
        total_llm_calls = 0
        llm_provider = "gemini"
        gemini_model = "gemini-3.5-flash"
        ollama_model = "gemma4:e4b"
        ollama_base_url = "http://localhost:11434"
        ollama_think = False
        try:
            # 2. Settings parameters
            top_k = int(get_setting('top_k', '3'))
            threshold = float(get_setting('similarity_threshold', '0.5'))
            chunk_size = int(get_setting('chunk_size', '50'))
            
            log_msg(f"[CONFIG] Parametry algorytmu: top_k={top_k}, threshold={threshold}, chunk_size={chunk_size}")
            
            # Divide input text into segments (smaller window sizes are better for search granularity)
            # We use 30 words window with 15 words overlap as a highly responsive sliding window
            segment_size = min(30, chunk_size)
            overlap = segment_size // 2
            
            segments = segment_text_with_indices(input_text, segment_size, overlap)
            log_msg(f"[SEGMENTACJA] Tekst podzielony na {len(segments)} segmentów (rozmiar segmentu: {segment_size} słów, overlap: {overlap} słów)")
            
            # Load LLM settings
            llm_provider = get_setting('llm_provider', 'gemini')
            gemini_model = get_setting('gemini_model', 'gemini-3.5-flash')
            ollama_model = get_setting('ollama_model', 'gemma4:e4b')
            ollama_base_url = get_setting('ollama_base_url', 'http://localhost:11434')
            ollama_think = setting_to_bool(get_setting('ollama_think', 'false'))
            temp = float(get_setting('gemini_temp', '0.1'))
            thinking_level = get_setting('gemini_thinking_level', 'low')
            active_model = ollama_model if llm_provider == "ollama" else gemini_model
            analysis_params = {
                "llm_provider": llm_provider,
                "llm_model": active_model,
                "embedding_model": "gemini-embedding-2",
                "gemini_temp": temp,
                "gemini_thinking_level": thinking_level if llm_provider == "gemini" else "",
                "ollama_base_url": ollama_base_url if llm_provider == "ollama" else "",
                "ollama_think": ollama_think if llm_provider == "ollama" else None,
                "top_k": top_k,
                "similarity_threshold": threshold,
                "corpus_chunk_size": chunk_size,
                "analysis_segment_size": segment_size,
                "analysis_segment_overlap": overlap,
                "segments_count": len(segments)
            }
            max_llm_validations = len(segments) * top_k
            if llm_provider == "ollama":
                log_msg(f"[LLM] Provider: Ollama, model: {ollama_model}, URL: {ollama_base_url}, temperatura: {temp}, thinking: {'włączone' if ollama_think else 'wyłączone'}")
            else:
                log_msg(f"[LLM] Provider: Gemini, model: {gemini_model}, temperatura: {temp}, poziom myślenia: {thinking_level}")
            log_msg(f"[KOSZTY] Maksymalna liczba walidacji LLM przed filtrem podobieństwa: {len(segments)} segmentów x top_k={top_k} = {max_llm_validations}")
            
            # 3. For each segment, search ChromaDB
            for idx, seg in enumerate(segments):
                seg_num = idx + 1
                seg_text_preview = seg['text'].replace('\n', ' ')
                if len(seg_text_preview) > 60:
                    seg_text_preview = seg_text_preview[:57] + "..."
                
                log_msg(f"[SEGMENT {seg_num}/{len(segments)}] Analiza fragmentu: \"{seg_text_preview}\"")
                
                # Zliczamy tokeny dla osadzenia zapytania (szacunkowo 1 token na 3 znaki łacińskie)
                total_embedding_tokens += len(seg['text']) // 3 + 1
                matches = vector_db.search_similar_chunks(seg['text'], top_k=top_k, similarity_threshold=threshold)
                log_msg(f"[ChromaDB] Szukanie podobnych fragmentów w bazie wektorowej. Znaleziono kandydatów spełniających próg podobieństwa: {len(matches)}")
                
                # If we have candidates, run LLM validation on each returned match.
                if matches:
                    for match_idx, match in enumerate(matches, start=1):
                        log_msg(
                            f"[ChromaDB] Kandydat {match_idx}/{len(matches)}: "
                            f"ID dok={match['corpus_doc_id']}, chunk={match['chunk_index']}, "
                            f"similarity={match['similarity']:.4f}"
                        )
                        
                        # Fetch adjacent chunks to reconstruct full context
                        expanded_corpus_text = vector_db.get_expanded_chunk_context(
                            match['corpus_doc_id'],
                            match['chunk_index']
                        )
                        
                        # Fallback to the single chunk if expanded retrieval fails
                        if not expanded_corpus_text:
                            expanded_corpus_text = match['text']
                            log_msg("[ChromaDB] Brak kontekstu rozszerzonego w bazie wektorowej - użycie podstawowego chunk'a")
                        else:
                            log_msg(f"[ChromaDB] Pobrano kontekst rozszerzony (długość: {len(expanded_corpus_text)} znaków)")
                        
                        try:
                            # Use expanded corpus text context from ChromaDB (preceding + target + succeeding chunks)
                            # to reduce API costs and improve prompt response times.
                            corpus_text = expanded_corpus_text
                            
                            # LLM Verification with expanded chunk context
                            log_msg(f"[LLM] Wywoływanie modelu {active_model} ({llm_provider}) w celu weryfikacji powiązania z tekstem...")
                            llm_result = embeddings.analyze_borrowing(
                                source_segment=seg['text'],
                                corpus_chunk=corpus_text,
                                model_name=active_model,
                                temperature=temp,
                                thinking_level=thinking_level,
                                provider=llm_provider,
                                ollama_base_url=ollama_base_url,
                                ollama_think=ollama_think
                            )
                            
                            # Zliczanie zużytych tokenów
                            usage = llm_result.get('usage')
                            prompt_delta, candidates_delta, thoughts_delta, calls_delta = add_usage_totals(
                                usage,
                                log_msg,
                                context=f"segment {seg_num}, kandydat {match_idx}"
                            )
                            total_prompt_tokens += prompt_delta
                            total_candidates_tokens += candidates_delta
                            total_thoughts_tokens += thoughts_delta
                            total_llm_calls += calls_delta
                            
                            is_match = llm_result.get('match')
                            match_type = llm_result.get('type', 'None')
                            log_msg(f"[LLM] Wynik analizy: zapożyczenie={is_match}, typ={match_type}")
                            
                            # If LLM confirms a borrowing relationship
                            if is_match and match_type != 'None':
                                borrowed_str = llm_result.get('borrowed_text_in_source', '').strip()
                                # Clean surrounding quotes
                                if borrowed_str.startswith('"') and borrowed_str.endswith('"'):
                                    borrowed_str = borrowed_str[1:-1].strip()
                                elif borrowed_str.startswith("'") and borrowed_str.endswith("'"):
                                    borrowed_str = borrowed_str[1:-1].strip()
                                    
                                # Pinpoint the exact character range in the analyzed segment
                                if borrowed_str and borrowed_str in seg['text']:
                                    relative_start = seg['text'].index(borrowed_str)
                                    start_char_idx = seg['start_char'] + relative_start
                                    end_char_idx = start_char_idx + len(borrowed_str)
                                    source_snippet = borrowed_str
                                    log_msg(f"[LLM] Dopasowano dokładną frazę w badanym tekście: \"{source_snippet[:40]}...\"")
                                else:
                                    start_char_idx = seg['start_char']
                                    end_char_idx = seg['end_char']
                                    source_snippet = seg['text']
                                    log_msg(f"[WARNING] Nie odnaleziono dokładnej frazy LLM '{borrowed_str[:30]}...' w badanym segmencie. Użycie pełnego segmentu.")
                                    
                                corpus_snippet = llm_result.get('borrowed_text_in_corpus', '').strip()
                                if corpus_snippet.startswith('"') and corpus_snippet.endswith('"'):
                                    corpus_snippet = corpus_snippet[1:-1].strip()
                                elif corpus_snippet.startswith("'") and corpus_snippet.endswith("'"):
                                    corpus_snippet = corpus_snippet[1:-1].strip()
                                    
                                if not corpus_snippet:
                                    corpus_snippet = expanded_corpus_text
                                    
                                log_msg(f"[SUCCESS] Wykryto powiązanie! Typ: {match_type}, Uzasadnienie: \"{llm_result.get('explanation', '')[:60]}...\"")
                                
                                # Create temporary detection in memory
                                detection = Detection(
                                    analysis_id=analysis_doc.id,                    # type: ignore
                                    source_snippet=source_snippet,                  # type: ignore
                                    corpus_document_id=match['corpus_doc_id'],      # type: ignore
                                    corpus_snippet_text=corpus_snippet,             # type: ignore
                                    match_type=match_type,                          # type: ignore 
                                    similarity_score=match['similarity'],           # type: ignore
                                    explanation=llm_result.get('explanation', ''),  # type: ignore
                                    start_char_idx=start_char_idx,                  # type: ignore  
                                    end_char_idx=end_char_idx                       # type: ignore 
                                )
                                detection._corpus_chunk_index = match['chunk_index']   # type: ignore
                                detection._corpus_chunk_start = match['chunk_index']   # type: ignore
                                detection._corpus_chunk_end = match['chunk_index']     # type: ignore
                                detection._source_fragments = [{                       # type: ignore
                                    "chunk_index": match['chunk_index'],
                                    "text": corpus_snippet,
                                    "similarity": match['similarity'],
                                    "match_type": match_type
                                }]
                                temp_detections.append(detection)

                                remaining_candidates = len(matches) - match_idx
                                if remaining_candidates > 0:
                                    log_msg(
                                        f"[LLM] Kandydat {match_idx} potwierdził zapożyczenie. "
                                        f"Pominięto {remaining_candidates} kolejnych kandydatów dla tego segmentu w celu ograniczenia kosztów."
                                    )
                                break
                            else:
                                no_match_explanation = (llm_result.get('explanation') or '').strip()
                                if no_match_explanation:
                                    log_msg(
                                        f"[INFO] Brak potwierdzenia zapożyczenia przez LLM "
                                        f"(wynik: 'None' lub brak dopasowania). Uzasadnienie: "
                                        f"\"{no_match_explanation}\""
                                    )
                                else:
                                    log_msg(f"[INFO] Brak potwierdzenia zapożyczenia przez LLM (wynik: 'None' lub brak dopasowania)")
                                
                        except Exception as llm_err:
                            err_str = str(llm_err)
                            log_msg(f"[ERROR] Błąd analizy LLM dla segmentu {seg_num}, kandydata {match_idx}: {err_str}")
                            log_msg(f"[ERROR] Traceback: {traceback.format_exc()}")
                            
                            # Sprawdzamy czy to błąd 429 (przekroczenie limitów/kosztów) lub błąd uwierzytelnienia klucza API, aby przerwać pętlę
                            is_fatal_api_error = any(kw in err_str.lower() or kw in type(llm_err).__name__.lower() 
                                                     for kw in ["429", "resource_exhausted", "quota", "limit", "api_key_invalid", "403", "unauthorized", "invalid_argument", "400", "ollama", "nie udało się połączyć"])
                            if is_fatal_api_error:
                                log_msg("[FATAL] Wykryto krytyczny błąd providera LLM. Przerywanie analizy w celu zachowania dotychczasowych wyników.")
                                raise ValueError("Krytyczny błąd providera LLM. Dla Gemini sprawdź limity/API key; dla Ollama sprawdź, czy serwer działa, adres jest poprawny i model został pobrany.") from llm_err
                            
                            pass
                else:
                    log_msg(f"[INFO] Pomiędzy segmentem {seg_num} a dokumentami w korpusie nie ma podobieństwa powyżej progu {threshold:.2f}")
            
            # Merge overlapping or close detections from the same source document
            log_msg(f"[SCALANIE] Przystępowanie do scalania nakładających się detekcji (liczba detekcji przed scaleniem: {len(temp_detections)})")
            merged_detections = merge_close_detections(temp_detections, input_text)
            log_msg(f"[SCALANIE] Zakończono scalanie. Liczba detekcji po scaleniu: {len(merged_detections)}")
            
            # Consolidate explanations of merged detections using LLM
            for idx_det, det in enumerate(merged_detections):
                if det.explanation and "\n[Ciąg dalszy]:" in det.explanation:
                    parts = [p.strip() for p in det.explanation.split("\n[Ciąg dalszy]:") if p.strip()]
                    if len(parts) > 1:
                        log_msg(f"[LLM] Wykryto scaloną detekcję (liczba segmentów: {len(parts)}). Rozpoczynanie konsolidacji wyjaśnień przez LLM...")
                        try:
                            consolidation_result = embeddings.consolidate_explanations(
                                explanations=parts,
                                model_name=active_model,
                                temperature=temp,
                                thinking_level=thinking_level,
                                provider=llm_provider,
                                ollama_base_url=ollama_base_url,
                                ollama_think=ollama_think
                            )
                            det.explanation = consolidation_result.get("text", det.explanation)
                            
                            usage = consolidation_result.get("usage")
                            prompt_delta, candidates_delta, thoughts_delta, calls_delta = add_usage_totals(
                                usage,
                                log_msg,
                                context="konsolidacja wyjaśnień"
                            )
                            total_prompt_tokens += prompt_delta
                            total_candidates_tokens += candidates_delta
                            total_thoughts_tokens += thoughts_delta
                            total_llm_calls += calls_delta
                                
                            log_msg(f"[LLM] Pomyślnie skonsolidowano wyjaśnienia.")
                        except Exception as cons_err:
                            log_msg(f"[WARNING] Błąd konsolidacji wyjaśnień przez LLM: {str(cons_err)}. Użycie połączonych tekstów.")
            
            estimated_cost = estimate_api_cost(
                total_prompt_tokens,
                total_candidates_tokens,
                total_thoughts_tokens,
                total_embedding_tokens,
                active_model,
                provider=llm_provider
            )
            
            log_msg(f"[KOSZTY] Liczba wywołań LLM: {total_llm_calls}")
            log_msg(f"[KOSZTY] Zużycie tokenów LLM: wejściowe={total_prompt_tokens}, wyjściowe={total_candidates_tokens}, myślowe={total_thoughts_tokens}")
            log_msg(f"[KOSZTY] Zużycie tokenów Embedding (zapytania, szacunek): {total_embedding_tokens}")
            if llm_provider == "ollama":
                log_msg(f"[KOSZTY] Szacowany koszt API: ${estimated_cost:.6f} USD (lokalny LLM przez Ollama bez kosztu API; koszt obejmuje embeddingi Gemini)")
            else:
                log_msg(f"[KOSZTY] Szacowany koszt Gemini API ({gemini_model} + embedding, output zawiera tokeny myślowe): ${estimated_cost:.6f} USD")
            
            # Save merged detections to SQLite
            for det in merged_detections:
                db.session.add(det)
                
            log_msg("Zapisywanie logów, detekcji i podsumowania kosztów do bazy SQLite...")
            analysis_doc.prompt_tokens = total_prompt_tokens
            analysis_doc.candidates_tokens = total_candidates_tokens
            analysis_doc.thoughts_tokens = total_thoughts_tokens
            analysis_doc.llm_calls = total_llm_calls
            analysis_doc.analysis_params = json.dumps(analysis_params, ensure_ascii=False)
            analysis_doc.estimated_cost = estimated_cost
            analysis_doc.logs = "\n".join(log_lines)
            db.session.commit()
            log_msg(f"Pomyślnie zakończono analizę. Zapisano {len(merged_detections)} detekcji.")
            
            # Save again to include the final success messages in the saved log
            analysis_doc.logs = "\n".join(log_lines)
            db.session.commit()
            
            return jsonify({
                "status": "success",
                "analysis_id": analysis_doc.id,
                "detections_count": len(merged_detections)
            })
            
        except Exception as e:
            db.session.rollback()
            error_trace = traceback.format_exc()
            log_msg(f"[FATAL] Krytyczny błąd w głównej pętli analizy: {str(e)}")
            log_msg(f"[FATAL] Traceback:\n{error_trace}")
            
            estimated_cost = estimate_api_cost(
                total_prompt_tokens,
                total_candidates_tokens,
                total_thoughts_tokens,
                total_embedding_tokens,
                active_model if 'active_model' in locals() else gemini_model,
                provider=llm_provider
            )
            
            log_msg(f"[KOSZTY] Awaryjne podsumowanie wywołań LLM: {total_llm_calls}")
            log_msg(f"[KOSZTY] Awaryjne podsumowanie tokenów: wejściowe={total_prompt_tokens}, wyjściowe={total_candidates_tokens}, myślowe={total_thoughts_tokens}, embedding={total_embedding_tokens}")
            log_msg(f"[KOSZTY] Awaryjny koszt zapytań: ${estimated_cost:.6f} USD")
            
            try:
                # Save partial logs and detections so the user can see what failed and keep partial results
                doc = AnalysisDocument.query.get(analysis_doc.id)
                if doc:
                    # Merge whatever detections we found up to now and save them
                    if temp_detections:
                        log_msg(f"[SCALANIE AWARYJNE] Przystępowanie do scalania detekcji wykrytych przed błędem (liczba przed: {len(temp_detections)})")
                        merged_detections = merge_close_detections(temp_detections, input_text)
                        for det in merged_detections:
                            db.session.add(det)
                        log_msg(f"[SCALANIE AWARYJNE] Zapisano awaryjnie {len(merged_detections)} detekcji.")
                    
                    doc.prompt_tokens = total_prompt_tokens
                    doc.candidates_tokens = total_candidates_tokens
                    doc.thoughts_tokens = total_thoughts_tokens
                    doc.llm_calls = total_llm_calls
                    doc.analysis_params = json.dumps(analysis_params, ensure_ascii=False) if 'analysis_params' in locals() else None
                    doc.estimated_cost = estimated_cost
                    doc.logs = "\n".join(log_lines)
                    db.session.commit()
                    log_msg("[INFO] Pomyślnie zapisano częściowe wyniki oraz logi analizy w bazie SQLite.")
            except Exception as db_err:
                print(f"Failed to save error logs/detections to database: {db_err}")
                db.session.rollback()
                
            return jsonify({
                "status": "error",
                "error": str(e),
                "analysis_id": analysis_doc.id,
                "trace": error_trace
            }), 500

@app.route('/api/analyze/<int:analysis_id>', methods=['GET', 'PUT', 'DELETE'])
def api_get_analysis(analysis_id):
    analysis = AnalysisDocument.query.get(analysis_id)
    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404
        
    if request.method == 'PUT':
        data = request.get_json() or {}
        title = data.get('title', '').strip()
        if not title:
            return jsonify({"error": "Title is required"}), 400
            
        try:
            analysis.title = title
            db.session.commit()
            return jsonify({"status": "success", "message": "Analysis title updated successfully"})
        except Exception as e:
            db.session.rollback()
            return jsonify({"error": str(e)}), 500
            
    elif request.method == 'GET':
        detections = []
        visible_detections = filter_visible_detections(analysis.detections)
        for det in visible_detections:
            detections.append({
                "id": det.id,
                "source_snippet": det.source_snippet,
                "corpus_document_id": det.corpus_document_id,
                "corpus_document_title": det.corpus_document.title,
                "corpus_document_author": det.corpus_document.author or "Unknown",
                "corpus_snippet_text": det.corpus_snippet_text,
                "corpus_source_fragments": parse_source_fragments(det),
                "match_type": det.match_type,
                "similarity_score": round(det.similarity_score, 4) if det.similarity_score else None,
                "explanation": det.explanation,
                "start_char_idx": det.start_char_idx,
                "end_char_idx": det.end_char_idx
            })
            
        return jsonify({
            "id": analysis.id,
            "title": analysis.title,
            "content": analysis.content,
            "logs": analysis.logs or "Brak logów dla tej analizy.",
            "prompt_tokens": analysis.prompt_tokens or 0,
            "candidates_tokens": analysis.candidates_tokens or 0,
            "thoughts_tokens": analysis.thoughts_tokens or 0,
            "llm_calls": analysis.llm_calls or 0,
            "analysis_params": parse_analysis_params(analysis),
            "estimated_cost": round(analysis.estimated_cost, 6) if analysis.estimated_cost else 0.0,
            "analyzed_at": analysis.analyzed_at.strftime("%Y-%m-%d %H:%M:%S"),
            "detections": detections
        })
        
    else:
        try:
            db.session.delete(analysis)
            db.session.commit()
            return jsonify({"status": "success", "message": f"Analysis {analysis_id} deleted successfully"})
        except Exception as e:
            db.session.rollback()
            return jsonify({"error": str(e)}), 500

@app.route('/api/analyze/<int:analysis_id>/export', methods=['GET'])
def api_export_analysis(analysis_id):
    analysis = AnalysisDocument.query.get(analysis_id)
    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404
        
    try:
        import io
        from docx import Document
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from flask import send_file
        
        doc = Document()
        
        # Heading
        doc.add_heading(analysis.title, level=0)
        
        # Metadata
        visible_detections = filter_visible_detections(analysis.detections)
        add_export_metadata(doc, analysis, detections_count=len(visible_detections))
            
        shading_color_map = {
            "Full": "FFC5C5",       # Jasny pastelowy czerwony (jak w aplikacji)
            "Partial": "FFE0B2",    # Jasny pastelowy pomarańczowy / bursztynowy
            "Paraphrase": "FFFDE7", # Jasny pastelowy żółty
            "Allusion": "EDE7F6"    # Jasny pastelowy fioletowy
        }
        
        content = analysis.content
        paragraphs_text = content.split('\n')
        
        p_starts = []
        current_offset = 0
        for p_text in paragraphs_text:
            p_starts.append(current_offset)
            current_offset += len(p_text) + 1  # +1 for newline character
            
        detections = visible_detections
        
        for idx, p_text in enumerate(paragraphs_text):
            p_start = p_starts[idx]
            p_end = p_start + len(p_text)
            
            paragraph = doc.add_paragraph()
            
            overlapping = []
            for det in detections:
                if det.start_char_idx < p_end and det.end_char_idx > p_start:
                    overlapping.append(det)
                    
            if not overlapping:
                paragraph.add_run(p_text)
            else:
                overlapping = sorted(overlapping, key=lambda d: d.start_char_idx)
                
                current_p_char = 0
                for det in overlapping:
                    rel_start = max(det.start_char_idx, p_start) - p_start
                    rel_end = min(det.end_char_idx, p_end) - p_start
                    
                    # Add plain text before highlighted run
                    if rel_start > current_p_char:
                        paragraph.add_run(p_text[current_p_char:rel_start])
                        
                    # Highlighted run (split to keep comment range markup overlay on the last character only)
                    h_text = p_text[rel_start:rel_end]
                    
                    if len(h_text) > 1:
                        h_run_main = paragraph.add_run(h_text[:-1])
                        h_run_anchor = paragraph.add_run(h_text[-1])
                    else:
                        h_run_main = None
                        h_run_anchor = paragraph.add_run(h_text)
                    
                    fill_color = shading_color_map.get(det.match_type)
                    if fill_color:
                        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>'
                        if h_run_main:
                            h_run_main._r.get_or_add_rPr().append(parse_xml(shading_xml))
                        h_run_anchor._r.get_or_add_rPr().append(parse_xml(shading_xml))
                        
                    source_fragments_text = "\n\n".join(
                        f"[chunk {fragment.get('chunk_index')}] {fragment.get('text')}"
                        if fragment.get('chunk_index') is not None
                        else fragment.get('text', '')
                        for fragment in parse_source_fragments(det)
                    )

                    # Detailed comment text
                    comment_text = (
                        f"Typ zapożyczenia: {det.match_type}\n"
                        f"Dokument źródłowy: {det.corpus_document.title}\n\n"
                        f"Pasujące fragmenty z korpusu:\n{source_fragments_text}\n\n"
                        f"Analiza semantyczna modelu:\n{det.explanation}"
                    )
                    
                    doc.add_comment(
                        runs=h_run_anchor,
                        text=comment_text,
                        author="Analiza zapożyczeń",
                        initials="AZ"
                    )
                    
                    current_p_char = rel_end
                    
                # Add plain text after the last highlighted run
                if current_p_char < len(p_text):
                    paragraph.add_run(p_text[current_p_char:])
                    
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = safe_docx_filename(analysis.title, "Analiza_zapożyczeń")
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        import traceback
        print(f"Failed to export analysis docx: {traceback.format_exc()}")
        return jsonify({"error": f"Failed to export: {str(e)}"}), 500

@app.route('/api/analyze/<int:analysis_id>/export-list', methods=['GET'])
def api_export_analysis_list(analysis_id):
    analysis = AnalysisDocument.query.get(analysis_id)
    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404

    try:
        import io
        from docx import Document
        from flask import send_file

        source_shading_color_map = {
            "Full": "FFC5C5",
            "Partial": "FFE0B2",
            "Paraphrase": "FFFDE7",
            "Allusion": "EDE7F6"
        }
        corpus_highlight_color = "FFF59D"

        doc = Document()
        doc.add_heading(f"{analysis.title} - lista zapożyczeń", level=0)

        detections = filter_visible_detections(analysis.detections)
        add_export_metadata(doc, analysis, detections_count=len(detections))
        if not detections:
            doc.add_paragraph("Nie wykryto zapożyczeń.")

        for idx, det in enumerate(detections, start=1):
            if idx > 1:
                add_horizontal_rule(doc)

            doc.add_heading(f"Zapożyczenie {idx}: {det.match_type}", level=1)

            doc.add_paragraph("Badany tekst", style=None).runs[0].bold = True
            before, highlighted, after = get_text_window_by_char(
                analysis.content,
                det.start_char_idx,
                det.end_char_idx,
                word_radius=10
            )
            add_highlighted_text_paragraph(
                doc,
                before,
                highlighted,
                after,
                source_shading_color_map.get(det.match_type, "FFFDE7")
            )

            source_p = doc.add_paragraph()
            source_p.add_run("Dokument referencyjny: ").bold = True
            source_p.add_run(det.corpus_document.title)
            if det.corpus_document.author:
                source_p.add_run(f" ({det.corpus_document.author})")
            if det.similarity_score is not None:
                source_p.add_run(f"\nPodobieństwo wektorowe: {det.similarity_score * 100:.1f}%")

            doc.add_paragraph("Pasujące fragmenty z korpusu", style=None).runs[0].bold = True
            fragments = parse_source_fragments(det)
            for fragment_idx, fragment in enumerate(fragments, start=1):
                chunk_index = fragment.get("chunk_index")
                fragment_text = (fragment.get("text") or "").strip()

                label_p = doc.add_paragraph()
                label_p.add_run(f"Fragment {fragment_idx}").bold = True
                if chunk_index is not None:
                    label_p.add_run(f" - chunk {chunk_index}")
                if fragment.get("similarity") is not None:
                    label_p.add_run(f" - podobieństwo: {fragment['similarity'] * 100:.1f}%")

                chunk_text = ""
                if chunk_index is not None:
                    chunk_text = vector_db.get_chunk_text(det.corpus_document_id, chunk_index)
                    if fragment_text and not text_contains_highlight(chunk_text, fragment_text):
                        expanded_chunk_text = vector_db.get_expanded_chunk_context(
                            det.corpus_document_id,
                            chunk_index
                        )
                        if text_contains_highlight(expanded_chunk_text, fragment_text):
                            chunk_text = expanded_chunk_text
                display_text = chunk_text or fragment_text

                if chunk_text and fragment_text:
                    before_corpus, highlighted_corpus, after_corpus = split_text_for_highlight(
                        chunk_text,
                        fragment_text
                    )
                    if not highlighted_corpus:
                        before_corpus, highlighted_corpus, after_corpus = "", fragment_text, ""
                else:
                    before_corpus, highlighted_corpus, after_corpus = "", display_text, ""

                add_highlighted_text_paragraph(
                    doc,
                    before_corpus,
                    highlighted_corpus,
                    after_corpus,
                    corpus_highlight_color
                )

            doc.add_paragraph("Analiza semantyczna modelu", style=None).runs[0].bold = True
            doc.add_paragraph(det.explanation or "Brak uzasadnienia.")

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = safe_docx_filename(analysis.title, "Lista_zapożyczeń")
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        import traceback
        print(f"Failed to export analysis list docx: {traceback.format_exc()}")
        return jsonify({"error": f"Failed to export list: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
